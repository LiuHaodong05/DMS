import docx
from docx.shared import Pt, RGBColor, Inches
import sys, io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = docx.Document(r'C:\Users\LiuHaodong\Desktop\WEB应用系统开发-课程设计文档(2).docx')
output_path = r'C:\Users\LiuHaodong\Desktop\WEB应用系统开发-课程设计文档(2).docx'

def set_run(run, text, font_name='宋体', font_size=12, bold=False, color=None):
    run.text = text
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# =============================================
# 1. 修改段落内容
# =============================================

# [18] 系统概述
p = doc.paragraphs[18]
p.clear()
set_run(p.add_run(), "随着高校招生规模的不断扩大，学生住宿管理工作日益复杂，传统的纸质登记和人工管理方式已难以满足现代高校宿舍管理的需求。本系统基于B/S架构，采用PHP语言结合MySQL数据库开发，旨在为高校提供一套高效、便捷的学生宿舍管理解决方案。系统实现了学生信息管理、宿舍分配、床位调度、数据统计等核心功能，能够有效提高宿舍管理的工作效率，降低管理成本，实现宿舍管理的信息化和规范化。")

# [28] 系统功能需求
p = doc.paragraphs[28]
p.clear()
set_run(p.add_run(), "本系统采用PHP语言开发，使用MySQL数据库存储数据，系统主要包括管理员认证、学生信息管理、宿舍分配管理、床位调度管理等模块。通过本系统，管理员可以进行登录认证，对学生信息进行增删改查操作，为学生分配宿舍和床位，查看空缺床位统计信息。系统支持自动分配空闲床位、调换宿舍时自动释放和占用床位、实时统计在宿学生数和空闲床位数等功能。系统的实现使高校宿舍管理更加方便、快捷，大大节省了管理成本。")

# [32] 系统结构设计
p = doc.paragraphs[32]
p.clear()
set_run(p.add_run(), "本系统采用分层架构设计，分为表现层（前端页面）、业务逻辑层（PHP处理脚本）和数据访问层（MySQL数据库）。表现层负责与管理员交互，展示学生信息和操作界面；业务逻辑层处理具体的业务请求，如学生添加、宿舍分配、床位调度等；数据访问层负责数据的存储和检索。系统的功能模块主要包括：（1）管理员登录认证模块——验证管理员身份，保护系统安全；（2）学生信息管理模块——实现学生基本信息的添加、修改、删除和查询；（3）宿舍分配管理模块——为学生分配宿舍和床位，支持自动查找空闲床位；（4）床位统计模块——实时统计各宿舍的空闲床位数和入住率。系统结构设计如下图所示。")

# [46] 数据库设计介绍
p = doc.paragraphs[46]
p.clear()
set_run(p.add_run(), "本系统采用MySQL数据库作为系统数据库，数据库名称为dorm_system，数据库的主要功能就是存储本系统中的所有数据，以便进行操作。下面介绍数据库中的各个数据表的详细设计信息。")

# 更新表名标题
p = doc.paragraphs[47]
p.clear()
set_run(p.add_run(), "（1）管理员信息表（表名：admin）", bold=True)

p = doc.paragraphs[49]
p.clear()
set_run(p.add_run(), "（2）学生信息表（表名：student）", bold=True)

p = doc.paragraphs[51]
p.clear()
set_run(p.add_run(), "（3）宿舍信息表（表名：dormitory）", bold=True)

# =============================================
# 2. 更新数据库表结构
# =============================================

# Table 0 - admin 表
t0 = doc.tables[0]
while len(t0.rows) > 1:
    row = t0.rows[-1]
    t0._tbl.remove(row._tr)

data_admin = [
    ('id', 'int(11)', '管理员表主键ID，自动增长'),
    ('username', 'varchar(50)', '管理员用户名，唯一标识'),
    ('password', 'varchar(255)', '管理员登录密码'),
]
for fld, tp, desc in data_admin:
    row = t0.add_row()
    row.cells[0].text = fld
    row.cells[1].text = tp
    row.cells[2].text = desc

# Table 1 - student 表
t1 = doc.tables[1]
while len(t1.rows) > 1:
    row = t1.rows[-1]
    t1._tbl.remove(row._tr)

data_student = [
    ('id', 'int(11)', '学生表主键ID，自动增长'),
    ('student_no', 'varchar(20)', '学号，唯一标识'),
    ('name', 'varchar(50)', '学生姓名'),
    ('class_name', 'varchar(50)', '所属班级'),
    ('gender', "enum('男','女')", '学生性别'),
    ('dorm_id', 'int(11)', '所属宿舍ID，关联dormitory表'),
    ('bed_no', 'int(11)', '床位编号（1-4号床）'),
]
for fld, tp, desc in data_student:
    row = t1.add_row()
    row.cells[0].text = fld
    row.cells[1].text = tp
    row.cells[2].text = desc

# Table 2 - dormitory 表
t2 = doc.tables[2]
while len(t2.rows) > 1:
    row = t2.rows[-1]
    t2._tbl.remove(row._tr)

data_dorm = [
    ('id', 'int(11)', '宿舍表主键ID，自动增长'),
    ('building_no', 'int(11)', '楼栋编号'),
    ('floor_no', 'int(11)', '楼层号'),
    ('room_no', 'int(11)', '房间号'),
    ('capacity', 'int(11)', '宿舍可容纳床位数，默认4'),
    ('available_beds', 'int(11)', '当前空闲床位数'),
]
for fld, tp, desc in data_dorm:
    row = t2.add_row()
    row.cells[0].text = fld
    row.cells[1].text = tp
    row.cells[2].text = desc

# =============================================
# 3. 更新文件说明表
# =============================================
t3 = doc.tables[3]
while len(t3.rows) > 1:
    row = t3.rows[-1]
    t3._tbl.remove(row._tr)

files_data = [
    ('config.php', '数据库配置文件，连接MySQL数据库'),
    ('login.php', '管理员登录验证页面'),
    ('logout.php', '管理员退出登录处理脚本'),
    ('index.php', '系统首页，展示所有学生住宿信息及统计'),
    ('add_student.php', '添加学生及分配宿舍床位功能页面'),
    ('edit_student.php', '修改学生信息和调换宿舍功能页面'),
    ('vacant_beds.php', '查看空缺床位统计和入住率信息'),
    ('style.css', '系统样式表，统一暗色奢华风格UI设计'),
    ('backup/', '项目备份文件目录'),
    ('test_db.php', '数据库连接测试脚本'),
]
for col1, col2 in files_data:
    row = t3.add_row()
    row.cells[0].text = col1
    row.cells[1].text = col2

# =============================================
# 4. 填充核心代码模块
# =============================================

# [58] 连接数据库代码
p = doc.paragraphs[58]
p.clear()
code = """<?php
$host = 'localhost';
$user = 'root';
$pass = 'root';
$dbname = 'dorm_system';

$conn = new mysqli($host, $user, $pass, $dbname);
if ($conn->connect_error) {
    die("连接失败: " . $conn->connect_error);
}
$conn->set_charset("utf8");
?>"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)

# [61] 添加学生
p = doc.paragraphs[61]
p.clear()
code = """// 检查学号是否已存在
$check = $conn->query("SELECT id FROM student WHERE student_no='$student_no'");
if ($check->num_rows > 0) {
    $error = "学号已存在！";
} else {
    // 获取宿舍信息及空闲床位
    $dormInfo = $conn->query("SELECT * FROM dormitory WHERE id=$dorm_id")->fetch_assoc();
    $usedBeds = $conn->query("SELECT bed_no FROM student WHERE dorm_id=$dorm_id");
    $used = [];
    while($b = $usedBeds->fetch_assoc()) $used[] = $b['bed_no'];
    $freeBed = 0;
    for($i=1; $i<=4; $i++) if(!in_array($i, $used)) { $freeBed = $i; break; }

    if($freeBed) {
        $conn->query("INSERT INTO student (student_no, name, class_name, gender, dorm_id, bed_no)
                      VALUES ('$student_no','$name','$class_name','$gender',$dorm_id,$freeBed)");
        $conn->query("UPDATE dormitory SET available_beds = available_beds - 1 WHERE id=$dorm_id");
    }
}"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：添加学生时，系统首先检查学号唯一性，然后自动查找所选宿舍的空闲床位（1-4号），分配后更新该宿舍的空闲床位数。")

# [63] 修改学生
p = doc.paragraphs[63]
p.clear()
code = """// 处理床位变更（换宿舍）
if ($new_dorm_id != $old_dorm_id && $old_dorm_id) {
    // 释放原宿舍床位
    $conn->query("UPDATE dormitory SET available_beds = available_beds + 1 WHERE id=$old_dorm_id");

    // 查找新宿舍空闲床位
    $usedBeds = $conn->query("SELECT bed_no FROM student WHERE dorm_id=$new_dorm_id");
    $used = [];
    while($b = $usedBeds->fetch_assoc()) $used[] = $b['bed_no'];
    $freeBed = 0;
    for($i=1;$i<=4;$i++) if(!in_array($i,$used)) { $freeBed = $i; break; }

    if($freeBed) {
        $conn->query("UPDATE student SET student_no='$student_no', name='$name',
                      dorm_id=$new_dorm_id, bed_no=$freeBed WHERE id=$id");
        $conn->query("UPDATE dormitory SET available_beds = available_beds - 1 WHERE id=$new_dorm_id");
    }
}"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：修改学生信息时，如果调换了宿舍，系统会自动释放原宿舍床位，在新宿舍中查找空闲床位并分配，同时更新两个宿舍的空闲床位数。")

# [65] 删除学生
p = doc.paragraphs[65]
p.clear()
code = """// 获取学生宿舍信息以更新空闲床位
$stu = $conn->query("SELECT dorm_id, bed_no FROM student WHERE id=$id")->fetch_assoc();
if ($stu && $stu['dorm_id']) {
    $conn->query("UPDATE dormitory SET available_beds = available_beds + 1 WHERE id=" . $stu['dorm_id']);
}
// 删除学生记录
$conn->query("DELETE FROM student WHERE id=$id");"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：删除学生时，系统先查询该学生的宿舍信息，然后释放其占用的床位，最后删除学生记录，确保宿舍空闲床位数据的准确性。")

# [67] 查询学生
p = doc.paragraphs[67]
p.clear()
code = """// 查询所有学生（关联宿舍信息）
$students = $conn->query("SELECT s.*, d.building_no, d.floor_no, d.room_no
                          FROM student s
                          LEFT JOIN dormitory d ON s.dorm_id = d.id
                          ORDER BY s.id DESC");

// 统计数据
$total = $conn->query("SELECT COUNT(*) as count FROM student")->fetch_assoc()['count'];
$vacant = $conn->query("SELECT SUM(available_beds) as vacant FROM dormitory")->fetch_assoc()['vacant'];"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：系统通过LEFT JOIN关联student表和dormitory表，在展示学生信息的同时显示其宿舍位置。首页顶部统计卡片实时显示在宿学生总数和总空闲床位数。")

# [70] 添加宿舍（空缺床位页面）
p = doc.paragraphs[70]
p.clear()
code = """// 查询有空缺床位的宿舍
$sql = "SELECT *, CONCAT(building_no,'号楼',floor_no,'层',room_no,'室') as room_name
        FROM dormitory WHERE available_beds > 0
        ORDER BY building_no, floor_no, room_no";
$result = $conn->query($sql);

// 计算每个宿舍的入住率
$occupied = $row['capacity'] - $row['available_beds'];
$rate = round(($occupied / $row['capacity']) * 100);"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：空缺床位页面列出所有有空位的宿舍并计算入住率。管理员可以直观地看到哪些宿舍还有空位，方便安排新生入住。")

# [72] 修改宿舍
p = doc.paragraphs[72]
p.clear()
code = """<!-- 在添加学生时自动选择有空位的宿舍 -->
<select name="dorm_id" required>
    <option value="">--请选择有空缺床位的宿舍--</option>
    <?php while($d = $dorms->fetch_assoc()): ?>
    <option value="<?= $d['id'] ?>">
        <?= $d['full_name'] ?> (空床<?= $d['available_beds'] ?>个)
    </option>
    <?php endwhile; ?>
</select>"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：在添加学生时，系统自动过滤出有空位的宿舍供管理员选择，下拉列表显示每间宿舍的空闲床位数，方便合理分配。")

# [74] 删除宿舍（实际是删除学生释放床位）
p = doc.paragraphs[74]
p.clear()
code = """// 首页删除学生时自动释放床位
if (isset($_GET['del_id'])) {
    $id = intval($_GET['del_id']);
    $stu = $conn->query("SELECT dorm_id, bed_no FROM student WHERE id=$id")->fetch_assoc();
    if ($stu && $stu['dorm_id']) {
        $conn->query("UPDATE dormitory SET available_beds = available_beds + 1 WHERE id=" . $stu['dorm_id']);
    }
    $conn->query("DELETE FROM student WHERE id=$id");
}"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：删除学生时自动释放其占用的床位资源，保证宿舍床位数据的完整性和准确性。")

# [76] 查询宿舍
p = doc.paragraphs[76]
p.clear()
code = """// 查看空缺床位统计
$total_vacant = $conn->query("SELECT SUM(available_beds) as total FROM dormitory")->fetch_assoc()['total'];
$total_rooms = $result->num_rows;
$total_capacity = $conn->query("SELECT SUM(capacity) as total FROM dormitory")->fetch_assoc()['total'];"""
set_run(p.add_run(), code, font_name='Courier New', font_size=9)
set_run(p.add_run(), "\n\n说明：空缺床位页面通过多条SQL查询实现多维度的床位统计，包括空闲床位总数、有空位宿舍数、总床位容量和总体入住率等关键数据。")

# =============================================
# 5. 课程设计总结
# =============================================
p = doc.paragraphs[88]
p.clear()
summary = """本课程设计完成了学生宿舍管理系统的基本功能，包括管理员登录认证、学生信息的增删改查、宿舍分配与床位调度、空缺床位统计等。通过本次课程设计，我深入实践了PHP+MySQL的Web应用开发流程，掌握了数据库设计、SQL查询优化、前后端交互等关键技术。

系统完成了以下主要功能：
（1）管理员认证：基于Session的登录验证，保障系统安全；
（2）学生管理：支持学生信息的添加、修改、删除和查询；
（3）宿舍分配：智能分配空闲床位，支持换宿舍时自动释放和重新分配；
（4）数据统计：实时统计在宿人数、空缺床位和入住率等关键指标。

存在的不足与改进方向：
（1）目前系统仅支持管理员单角色，未来可增加学生端，让学生可以查看自己的住宿信息；
（2）密码存储目前为明文/MD5方式，应改用更安全的密码哈希算法（如bcrypt）；
（3）SQL语句存在注入风险，后续应改用参数化查询或ORM框架；
（4）移动端适配还有优化空间，可进一步完善响应式布局；
（5）可增加数据导出功能，如导出Excel格式的学生住宿名册。

通过本次课程设计，我不仅巩固了PHP和MySQL的编程知识，还锻炼了独立分析和解决问题的能力，为今后开发更复杂的Web应用打下了坚实基础。"""
set_run(p.add_run(), summary)

# =============================================
# 保存
# =============================================
doc.save(r'D:\phpstudy_pro\WWW\DMS\WEB应用系统开发-课程设计文档(完成版).docx')
print("文档已成功保存！")
print("路径: D:\\phpstudy_pro\\WWW\\DMS\\WEB应用系统开发-课程设计文档(完成版).docx")
